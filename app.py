import streamlit as st
import base64
import io
from cyber_yantra import encrypt_text, decrypt_text

# Page configuration
st.set_page_config(
    page_title="Cyber Yantra - Encryption Tool",
    page_icon="🔐",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Title and introduction
st.title("Cyber Yantra")
st.subheader("A Substitution Cipher-Based Encryption Tool for Secure Digital Communication")

# Create columns for layout
col1, col2 = st.columns([2, 1])

with col1:
    st.markdown("""
    ## Welcome to Cyber Yantra
    
    In the digital age, safeguarding sensitive information is paramount. Cyber Yantra is a 
    lightweight encryption tool employing a substitution cipher mechanism for basic text 
    confidentiality.
    
    This application demonstrates the concepts, research, and practical applications of the 
    Cyber Yantra encryption tool as described in the research paper by Tanmay S Dikshit, 
    Om Raju Golesar, Niraj Digambar Shinde, and Kishor M Mahale.
    
    ### What is Cyber Yantra?
    
    Cyber Yantra bridges classical and contemporary cryptographic approaches by leveraging 
    dynamic key-based substitutions combined with modular arithmetic. It offers a balanced 
    solution that is both accessible and secure, making it particularly suitable for 
    educational purposes and resource-constrained environments.
    """)

with col2:
    # Display a cybersecurity-related image (using a placeholder icon)
    st.markdown("""
    <div style="text-align: center; padding: 20px;">
        <i class="material-icons" style="font-size: 150px; color: #0066cc;">lock</i>
        <p>Secure your communications with Cyber Yantra</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Try the encryption tool directly
    st.markdown("### Try it now!")
    demo_text = st.text_input("Enter text to encrypt:", "Hello, Cyber Yantra!")
    key = st.slider("Select encryption key (1-25):", 1, 25, 7)
    
    if demo_text:
        encrypted = encrypt_text(demo_text, key)
        st.success(f"Encrypted: **{encrypted}**")
        
        if st.button("Decrypt"):
            decrypted = decrypt_text(encrypted, key)
            st.info(f"Decrypted: **{decrypted}**")

# Download research paper option
st.sidebar.title("Navigation")
st.sidebar.info("""
Use the navigation menu above to explore different sections:
1. Home (current page)
2. Research Overview
3. Importance of Secure Communication
4. Cyber Yantra Encryption Mechanism
5. Case Studies and Real-Life Applications
6. Interactive Data Visualizations
7. Conclusion & Future Scope
""")

# Load paper content
with open("assets/cyber_yantra_paper.txt", "r") as f:
    paper_content = f.read()

def get_download_link(text, filename, link_text, mime_type="text/plain"):
    """Generate a download link for text content"""
    b64 = base64.b64encode(text.encode()).decode()
    href = f'<a href="data:file/{mime_type};base64,{b64}" download="{filename}">{link_text}</a>'
    return href

def get_docx_download_link(original_content, filename, link_text):
    """Generate a download link for the research paper in DOCX format"""
    try:
        from docx import Document
        
        # Create a new Document
        document = Document()
        
        # Add title
        document.add_heading('Cyber Yantra: A Substitution Cipher-Based Encryption Tool for Secure Digital Communication', level=0)
        
        # Add authors section
        authors = "Tanmay S Dikshit, Om Raju Golesar, Niraj Digambar Shinde, Kishor M Mahale"
        document.add_paragraph(authors)
        document.add_paragraph('')  # Spacing
        
        # Add content paragraphs
        paragraphs = original_content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                if para.startswith('##') or para.startswith('# '):
                    # Handle headings
                    heading_text = para.lstrip('#').strip()
                    document.add_heading(heading_text, level=1)
                else:
                    document.add_paragraph(para)
        
        # Save to memory buffer
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        # Create base64 string
        b64 = base64.b64encode(buffer.read()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">{link_text}</a>'
        return href
        
    except ImportError:
        # If python-docx is not available, return a plain text download link instead
        return get_download_link(original_content, filename.replace('.docx', '.txt'), 
                              f"{link_text} (TXT format)", "text/plain")

def get_pdf_download_link(original_content, filename, link_text):
    """Generate a download link for the research paper in PDF format"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # Create buffer for PDF
        buffer = io.BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(buffer, pagesize=letter, 
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=72)
        
        # Get styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='Title', 
                                 parent=styles['Heading1'], 
                                 fontSize=16, 
                                 alignment=TA_CENTER))
        styles.add(ParagraphStyle(name='Authors', 
                                 parent=styles['Normal'], 
                                 fontSize=12, 
                                 alignment=TA_CENTER))
        styles.add(ParagraphStyle(name='Body', 
                                 parent=styles['Normal'], 
                                 fontSize=10, 
                                 alignment=TA_JUSTIFY))
        
        # Create content flow
        content = []
        
        # Add title
        title = "Cyber Yantra: A Substitution Cipher-Based Encryption Tool for Secure Digital Communication"
        content.append(Paragraph(title, styles['Title']))
        content.append(Spacer(1, 12))
        
        # Add authors
        authors = "Tanmay S Dikshit, Om Raju Golesar, Niraj Digambar Shinde, Kishor M Mahale"
        content.append(Paragraph(authors, styles['Authors']))
        content.append(Spacer(1, 24))
        
        # Process content paragraphs
        paragraphs = original_content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                if para.startswith('##') or para.startswith('# '):
                    # Handle headings
                    heading_text = para.lstrip('#').strip()
                    content.append(Spacer(1, 12))
                    content.append(Paragraph(heading_text, styles['Heading2']))
                    content.append(Spacer(1, 6))
                else:
                    content.append(Paragraph(para, styles['Body']))
                    content.append(Spacer(1, 10))
        
        # Build PDF
        doc.build(content)
        buffer.seek(0)
        
        # Create base64 string
        b64 = base64.b64encode(buffer.getvalue()).decode()
        href = f'<a href="data:application/pdf;base64,{b64}" download="{filename}">{link_text}</a>'
        return href
        
    except ImportError:
        # If reportlab is not available, return a plain text download link instead
        return get_download_link(original_content, filename.replace('.pdf', '.txt'), 
                              f"{link_text} (TXT format)", "text/plain")

st.sidebar.markdown("### Research Paper")
st.sidebar.markdown(get_download_link(paper_content, "Cyber_Yantra_Research_Paper.txt", "Download as TXT"), unsafe_allow_html=True)

# Add DOCX download option
try:
    docx_link = get_docx_download_link(paper_content, "Cyber_Yantra_Research_Paper.docx", "Download as DOCX")
    st.sidebar.markdown(docx_link, unsafe_allow_html=True)
except Exception as e:
    st.sidebar.warning(f"DOCX format not available: {str(e)}")

# Add PDF download option
try:
    pdf_link = get_pdf_download_link(paper_content, "Cyber_Yantra_Research_Paper.pdf", "Download as PDF")
    st.sidebar.markdown(pdf_link, unsafe_allow_html=True)
except Exception as e:
    st.sidebar.warning(f"PDF format not available: {str(e)}")

st.sidebar.markdown("---")

# Footer with credits
st.markdown("---")
st.markdown("""
<div style="text-align: center;">
    <small>
        Created based on the research by Tanmay S Dikshit, Om Raju Golesar, Niraj Digambar Shinde, Kishor M Mahale, Ratan D Deokar.<br>
        This application is for educational purposes only.
    </small>
</div>
""", unsafe_allow_html=True)


